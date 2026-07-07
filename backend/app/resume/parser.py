"""Resume file parsing: docx / pdf / plain text → raw text + detected skills."""
import io
import re

from app.core.errors import Problem

_SKILL_VOCAB = [
    "Python", "Java", "JavaScript", "TypeScript", "Go", "Rust", "C++", "C#",
    "Ruby", "Scala", "Kotlin", "SQL", "PostgreSQL", "MySQL", "MongoDB",
    "Redis", "Kafka", "Spark", "Airflow", "Snowflake", "AWS", "Azure", "GCP",
    "Docker", "Kubernetes", "Terraform", "React", "Angular", "Vue", "Node.js",
    "Spring Boot", "Spring", "Django", "FastAPI", "Flask", "GraphQL", "REST",
    "gRPC", "Microservices", "CI/CD", "Machine Learning", "TensorFlow",
    "PyTorch", "Pandas", "Jenkins", "Git", "Hibernate", "JPA", "Oracle",
]


def parse_file(filename: str, content: bytes, mime_type: str) -> str:
    """Extract plain text from an uploaded resume file."""
    lowered = (filename or "").lower()
    try:
        if lowered.endswith(".docx") or "wordprocessingml" in (mime_type or ""):
            from docx import Document
            doc = Document(io.BytesIO(content))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    parts.extend(c.text for c in row.cells if c.text.strip())
            return "\n".join(parts)
        if lowered.endswith(".pdf") or "pdf" in (mime_type or ""):
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(content))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        return content.decode("utf-8", errors="replace")
    except Exception as e:  # noqa: BLE001
        raise Problem(422, "Could not parse resume file",
                      f"Supported formats: docx, pdf, txt. ({str(e)[:100]})",
                      type_suffix="resume-parse") from e


def detect_skills(text: str) -> list[str]:
    return [s for s in _SKILL_VOCAB
            if re.search(rf"(?<![\w+#]){re.escape(s)}(?![\w+#])", text, re.I)]
