"""Phase 9 integration tests: resumes, tracker, match scoring pipeline,
and the analytics dashboard — full flow against real PostgreSQL."""
import io

from tests.test_auth import auth_headers, register


def _upload_resume(client, headers, name="Main resume") -> dict:
    from docx import Document
    buf = io.BytesIO()
    doc = Document()
    doc.add_paragraph("Rakesh Naini — rakesh@example.com — 469-813-5441")
    doc.add_paragraph("Skills: Java, Spring Boot, AWS, SQL, Kafka, Microservices")
    doc.add_paragraph("Experience: 6 years building backend microservices on "
                      "AWS with Java and Spring Boot, Kafka pipelines and SQL.")
    doc.save(buf)
    buf.seek(0)
    r = client.post(f"/api/v1/resumes?name={name}",
                    files={"file": ("resume.docx", buf.getvalue(),
                                    "application/vnd.openxmlformats-officedocument"
                                    ".wordprocessingml.document")},
                    headers=headers)
    assert r.status_code == 201, r.text
    return r.json()


def _seed_job(title="Senior Java Developer", company="Acme",
              description="Java, Spring Boot, AWS, Kafka. H1B sponsorship available.",
              url="https://acme.io/j1", salary_max=150000):
    from app.connector.spi import RawPosting
    from app.core.db import worker_session
    from app.ingestion.normalize import upsert_job
    with worker_session() as s:
        job, _ = upsert_job(s, RawPosting(
            connector_id="dice", title=title, company_name=company, url=url,
            description_md=description, salary_max=salary_max,
            posted_at="2026-07-06T00:00:00Z"))
        # backfill skills onto the extraction (normally done by AI enrichment)
        from app.connector.models import JobExtraction
        ext = s.get(JobExtraction, job.id)
        ext.skills = ["Java", "Spring Boot", "AWS", "Kafka", "Kubernetes"]
        return str(job.id)


# ---- resumes ---------------------------------------------------------------
def test_resume_upload_parse_and_analysis(client):
    pair = register(client).json()
    headers = auth_headers(pair)
    resume = _upload_resume(client, headers)
    assert resume["is_default"] is True                    # first upload → default
    assert "Java" in resume["structured"]["skills"]

    analysis = client.get(f"/api/v1/resumes/{resume['id']}/analysis",
                          headers=headers).json()
    assert analysis["ats_score"] > 0
    assert analysis["keyword_coverage"]["has_email"] is True


# ---- tracker ---------------------------------------------------------------
def test_tracker_crud_and_status_history(client):
    pair = register(client).json()
    headers = auth_headers(pair)
    job_id = _seed_job()

    created = client.post("/api/v1/applications",
                          json={"job_id": job_id}, headers=headers)
    assert created.status_code == 201
    app_id = created.json()["id"]
    assert created.json()["status"] == "SAVED"

    dup = client.post("/api/v1/applications", json={"job_id": job_id},
                      headers=headers)
    assert dup.status_code == 409                          # unique per user+job

    for status in ("APPLIED", "INTERVIEW_SCHEDULED", "OFFER"):
        r = client.post(f"/api/v1/applications/{app_id}/status",
                        json={"to_status": status}, headers=headers)
        assert r.status_code == 200
    assert r.json()["applied_at"] is not None              # stamped on APPLIED

    events = client.get(f"/api/v1/applications/{app_id}/events",
                        headers=headers).json()
    assert [e["to_status"] for e in events] == \
        ["SAVED", "APPLIED", "INTERVIEW_SCHEDULED", "OFFER"]

    listed = client.get("/api/v1/applications?status=OFFER", headers=headers).json()
    assert listed["total"] == 1
    assert listed["items"][0]["job"]["title"] == "Senior Java Developer"

    assert client.delete(f"/api/v1/applications/{app_id}",
                         headers=headers).status_code == 204
    assert client.get(f"/api/v1/applications/{app_id}",
                      headers=headers).status_code == 404  # soft-deleted


# ---- matching pipeline -----------------------------------------------------
def test_scoring_persists_and_surfaces_in_matches(client):
    pair = register(client).json()
    headers = auth_headers(pair)
    _upload_resume(client, headers)
    job_id = _seed_job()

    from app.core.db import worker_session
    from app.matching.service import score_jobs_for_user
    with worker_session() as s:
        import uuid
        scored = score_jobs_for_user(s, uuid.UUID(pair["user"]["id"]))
    assert scored == 1

    matches = client.get("/api/v1/matches", headers=headers).json()
    assert matches["total"] == 1
    m = matches["items"][0]
    assert m["job"]["title"] == "Senior Java Developer"
    assert m["score"]["overall"] > 50                      # 4/5 skills present
    missing = [g["skill"] for g in m["score"]["skill_gap"] if not g["have"]]
    assert missing == ["Kubernetes"]

    breakdown = client.get(f"/api/v1/jobs/{job_id}/match", headers=headers)
    assert breakdown.status_code == 200
    assert breakdown.json()["resume_pct"] == 80.0

    # scoring is idempotent (upsert, not duplicate)
    with worker_session() as s:
        score_jobs_for_user(s, uuid.UUID(pair["user"]["id"]))
    assert client.get("/api/v1/matches", headers=headers).json()["total"] == 1


def test_scoring_skipped_without_resume(client):
    pair = register(client).json()
    _seed_job()
    import uuid

    from app.core.db import worker_session
    from app.matching.service import score_jobs_for_user
    with worker_session() as s:
        assert score_jobs_for_user(s, uuid.UUID(pair["user"]["id"])) == 0


# ---- analytics dashboard ----------------------------------------------------
def test_dashboard_reflects_pipeline(client):
    pair = register(client).json()
    headers = auth_headers(pair)
    _upload_resume(client, headers)
    j1 = _seed_job(url="https://acme.io/a")
    j2 = _seed_job(title="Java Backend Engineer", url="https://acme.io/b",
                   description="Java and SQL.", salary_max=120000)

    import uuid

    from app.core.db import worker_session
    from app.matching.service import score_jobs_for_user
    with worker_session() as s:
        score_jobs_for_user(s, uuid.UUID(pair["user"]["id"]))

    a1 = client.post("/api/v1/applications", json={"job_id": j1, "status": "APPLIED"},
                     headers=headers).json()
    client.post("/api/v1/applications", json={"job_id": j2}, headers=headers)
    client.post(f"/api/v1/applications/{a1['id']}/status",
                json={"to_status": "INTERVIEW_SCHEDULED"}, headers=headers)

    d = client.get("/api/v1/analytics/dashboard", headers=headers).json()
    assert d["jobs_found"] == 2
    assert d["jobs_applied"] == 1
    assert d["interviews"] == 1
    assert d["recruiter_response_rate"] == 100.0
    assert len(d["top_matches"]) == 2
    assert any(f["status"] == "SAVED" for f in d["funnel"])
    assert d["match_score_distribution"]                    # buckets exist

    t = client.get("/api/v1/analytics/trends", headers=headers).json()
    assert t["postings_per_day"]
    assert t["tech_demand"][0]["skill"] in ("Java", "Spring Boot", "AWS", "Kafka",
                                            "Kubernetes")
    assert t["booming_companies"][0]["company"] == "Acme"
